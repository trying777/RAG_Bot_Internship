"""
Document Processing Module
Handles loading and chunking of different file types (PDF, TXT, MD, DOCX, CSV)
Includes OCR support for scanned PDFs and images
"""

import os
import PyPDF2
import markdown
from typing import List, Dict, Any
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain.schema import Document
from app.config import config

# OCR imports
try:
    import pytesseract
    import cv2
    import numpy as np
    from PIL import Image
    from pdf2image import convert_from_path
    OCR_AVAILABLE = True
except ImportError:
    OCR_AVAILABLE = False
    print("⚠️ OCR libraries not available. Install with: pip install pytesseract opencv-python pdf2image Pillow")

class DocumentProcessor:
    """Handles document loading and text chunking with OCR support"""
    
    def __init__(self):
        """Initialize document processor with text splitter and OCR setup"""
        self.text_splitter = RecursiveCharacterTextSplitter(
            chunk_size=config.CHUNK_SIZE,
            chunk_overlap=config.CHUNK_OVERLAP,
            length_function=len,
            separators=["\n\n", "\n", " ", ""]
        )
        
        # Setup OCR if available
        if OCR_AVAILABLE and config.ENABLE_OCR:
            self._setup_ocr()
        
        # Supported file extensions
        self.supported_extensions = {
            '.pdf': self._load_pdf,
            '.txt': self._load_txt,
            '.md': self._load_markdown,
            '.docx': self._load_docx,
            '.csv': self._load_csv,
            '.jpg': self._load_image,
            '.jpeg': self._load_image,
            '.png': self._load_image,
            '.bmp': self._load_image,
            '.tiff': self._load_image
        }
    
    def _setup_ocr(self):
        """Setup OCR configuration"""
        try:
            # Set Tesseract command path
            pytesseract.pytesseract.tesseract_cmd = config.TESSERACT_CMD_PATH
            
            # Set TESSDATA_PREFIX environment variable
            os.environ["TESSDATA_PREFIX"] = config.TESSDATA_PREFIX
            
            print(f"✓ OCR configured with Tesseract at: {config.TESSERACT_CMD_PATH}")
            print(f"✓ TESSDATA_PREFIX set to: {config.TESSDATA_PREFIX}")
            
        except Exception as e:
            print(f"⚠️ OCR setup warning: {str(e)}")
            print("  OCR will still work but may use system defaults")
    
    def _is_scanned_pdf(self, file_path: str) -> bool:
        """
        Detect if PDF is scanned (image-based) or text-based
        
        Args:
            file_path: Path to PDF file
            
        Returns:
            True if scanned PDF, False if text-based
        """
        try:
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                
                # Check first few pages for text content
                text_content = ""
                pages_to_check = min(3, len(pdf_reader.pages))
                
                for i in range(pages_to_check):
                    page = pdf_reader.pages[i]
                    text_content += page.extract_text() or ""
                
                # If very little text extracted, likely scanned
                if len(text_content.strip()) < 100:
                    print(f"🔍 Detected scanned PDF: {os.path.basename(file_path)}")
                    return True
                else:
                    print(f"🔍 Detected text-based PDF: {os.path.basename(file_path)}")
                    return False
                    
        except Exception as e:
            print(f"⚠️ Error detecting PDF type: {str(e)}")
            # Default to OCR processing for safety
            return True
    
    def _extract_text_with_ocr(self, image_path: str) -> str:
        """
        Extract text from image using OCR
        
        Args:
            image_path: Path to image file
            
        Returns:
            Extracted text
        """
        try:
            # Read image
            image = cv2.imread(image_path)
            if image is None:
                raise ValueError("Could not read image")
            
            # Convert to RGB (OpenCV uses BGR)
            image_rgb = cv2.cvtColor(image, cv2.COLOR_BGR2RGB)
            
            # Preprocess image for better OCR
            # Convert to grayscale
            gray = cv2.cvtColor(image_rgb, cv2.COLOR_RGB2GRAY)
            
            # Apply thresholding to get binary image
            _, binary = cv2.threshold(gray, 0, 255, cv2.THRESH_BINARY + cv2.THRESH_OTSU)
            
            # Extract text using Tesseract
            text = pytesseract.image_to_string(
                binary, 
                lang=config.OCR_LANGUAGE,
                config=config.OCR_CONFIG
            )
            
            print(f"✓ OCR extracted {len(text)} characters from image")
            return text
            
        except Exception as e:
            print(f"✗ OCR error: {str(e)}")
            return ""
    
    def _extract_text_from_scanned_pdf(self, file_path: str) -> str:
        """
        Extract text from scanned PDF using OCR
        
        Args:
            file_path: Path to scanned PDF file
            
        Returns:
            Extracted text
        """
        try:
            print(f"🔄 Processing scanned PDF with OCR: {os.path.basename(file_path)}")
            
            # Convert PDF pages to images
            pages = convert_from_path(file_path, dpi=300)
            
            extracted_text = ""
            
            for i, page in enumerate(pages):
                print(f"  Processing page {i+1}/{len(pages)}...")
                
                # Convert PIL image to OpenCV format
                page_cv = cv2.cvtColor(np.array(page), cv2.COLOR_RGB2BGR)
                
                # Convert to grayscale
                gray = cv2.cvtColor(page_cv, cv2.COLOR_BGR2GRAY)
                
                # Apply thresholding
                _, binary = cv2.threshold(gray, 0, 255, cv2.THRESH_BINARY + cv2.THRESH_OTSU)
                
                # Extract text from page
                page_text = pytesseract.image_to_string(
                    binary,
                    lang=config.OCR_LANGUAGE,
                    config=config.OCR_CONFIG
                )
                
                extracted_text += f"\n\n--- Page {i+1} ---\n\n{page_text}"
            
            print(f"✓ OCR extracted {len(extracted_text)} characters from {len(pages)} pages")
            return extracted_text
            
        except Exception as e:
            print(f"✗ Error processing scanned PDF: {str(e)}")
            return ""
    
    def load_document(self, file_path: str) -> str:
        """
        Load document content based on file extension
        
        Args:
            file_path: Path to the document file
            
        Returns:
            Extracted text content
        """
        file_extension = os.path.splitext(file_path)[1].lower()
        
        if file_extension not in self.supported_extensions:
            raise ValueError(f"Unsupported file type: {file_extension}")
        
        try:
            content = self.supported_extensions[file_extension](file_path)
            print(f"✓ Successfully loaded: {os.path.basename(file_path)}")
            return content
        except Exception as e:
            print(f"✗ Error loading {file_path}: {str(e)}")
            raise
    
    def _load_pdf(self, file_path: str) -> str:
        """Extract text from PDF file with OCR support for scanned PDFs"""
        # First check if it's a scanned PDF
        if self._is_scanned_pdf(file_path):
            if OCR_AVAILABLE and config.ENABLE_OCR:
                return self._extract_text_from_scanned_pdf(file_path)
            else:
                raise Exception("Scanned PDF detected but OCR is not available or enabled")
        else:
            # Text-based PDF - use standard extraction
            text = ""
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                for page in pdf_reader.pages:
                    text += page.extract_text() + "\n"
            return text
    
    def _load_txt(self, file_path: str) -> str:
        """Load text from TXT file"""
        with open(file_path, 'r', encoding='utf-8') as file:
            return file.read()
    
    def _load_markdown(self, file_path: str) -> str:
        """Load and convert markdown to text"""
        with open(file_path, 'r', encoding='utf-8') as file:
            md_content = file.read()
            # Convert markdown to HTML, then extract text
            html = markdown.markdown(md_content)
            # Simple HTML tag removal (for now)
            import re
            text = re.sub(r'<[^>]+>', '', html)
            return text
    
    def _load_docx(self, file_path: str) -> str:
        """Load text from DOCX file"""
        try:
            from docx import Document
            doc = Document(file_path)
            text = ""
            for paragraph in doc.paragraphs:
                text += paragraph.text + "\n"
            return text
        except ImportError:
            raise ImportError("python-docx not installed. Install with: pip install python-docx")
    
    def _load_csv(self, file_path: str) -> str:
        """Load text from CSV file"""
        import pandas as pd
        df = pd.read_csv(file_path)
        return df.to_string()
    
    def _load_image(self, file_path: str) -> str:
        """Extract text from image using OCR"""
        if not OCR_AVAILABLE or not config.ENABLE_OCR:
            raise Exception("OCR is not available or enabled for image processing")
        
        return self._extract_text_with_ocr(file_path)
    
    def chunk_document(self, content: str, metadata: Dict[str, Any] = None) -> List[Document]:
        """
        Split document content into chunks
        
        Args:
            content: Document text content
            metadata: Additional metadata for the document
            
        Returns:
            List of Document chunks
        """
        if metadata is None:
            metadata = {}
        
        # Create initial document
        doc = Document(page_content=content, metadata=metadata)
        
        # Split into chunks
        chunks = self.text_splitter.split_documents([doc])
        
        # Limit chunks if specified
        if config.MAX_CHUNKS_PER_DOCUMENT and len(chunks) > config.MAX_CHUNKS_PER_DOCUMENT:
            chunks = chunks[:config.MAX_CHUNKS_PER_DOCUMENT]
            print(f"⚠ Limited to {config.MAX_CHUNKS_PER_DOCUMENT} chunks")
        
        print(f"✓ Created {len(chunks)} chunks from document")
        return chunks
    
    def process_document(self, file_path: str) -> List[Document]:
        """
        Complete document processing pipeline
        
        Args:
            file_path: Path to the document file
            
        Returns:
            List of processed Document chunks
        """
        # Load document content
        content = self.load_document(file_path)
        
        # Prepare metadata
        metadata = {
            'source': file_path,
            'filename': os.path.basename(file_path),
            'file_type': os.path.splitext(file_path)[1].lower(),
            'file_size': os.path.getsize(file_path)
        }
        
        # Chunk the document
        chunks = self.chunk_document(content, metadata)
        
        return chunks
